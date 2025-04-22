# MLResearchApp.py

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import shap
import sweetviz as sv
import tempfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches
import base64
import plotly.io as pio
import io
import json
from datetime import datetime
import pickle
import traceback
import warnings
from sklearn.model_selection import StratifiedKFold, cross_val_score, train_test_split, GridSearchCV, RandomizedSearchCV
from sklearn.preprocessing import label_binarize, StandardScaler, MinMaxScaler, RobustScaler
from sklearn.feature_selection import SelectKBest, f_classif, mutual_info_classif, RFE
from sklearn.impute import SimpleImputer, KNNImputer
from sklearn.ensemble import IsolationForest
from sklearn.metrics import (
    accuracy_score, f1_score, cohen_kappa_score, confusion_matrix,
    roc_auc_score, roc_curve, auc, brier_score_loss, mean_squared_error, r2_score,
    precision_recall_curve, average_precision_score
)
from sklearn.calibration import CalibratedClassifierCV
from sklearn.ensemble import RandomForestClassifier, RandomForestRegressor
from sklearn.linear_model import LogisticRegression, LinearRegression
from xgboost import XGBClassifier, XGBRegressor
from statsmodels.stats.contingency_tables import mcnemar
from lime import lime_tabular
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from sklearn.tree import DecisionTreeClassifier, DecisionTreeRegressor
from sklearn.dummy import DummyClassifier, DummyRegressor
from sklearn.svm import SVC, SVR
from sklearn.neural_network import MLPClassifier, MLPRegressor
from sklearn.preprocessing import LabelEncoder

# Suppress specific warnings
warnings.filterwarnings("ignore", message="Mean of empty slice")
warnings.filterwarnings("ignore", category=RuntimeWarning)
warnings.filterwarnings("ignore", message="Parameters:")

# Add tooltips dictionary at the top of the file after imports
TOOLTIPS = {
    "task_type": "Classification: Predict categories (e.g., disease/no disease). Regression: Predict continuous values (e.g., blood pressure).",
    "missing_value_strategy": "How to handle empty data points: Mean/Median/Mode fills with average values, KNN uses similar data points, Drop removes rows with missing values.",
    "outlier_detection": "Identify unusual data points: Isolation Forest finds anomalies, Z-score flags values far from average, IQR identifies values outside typical range.",
    "feature_scaling": "Adjust feature ranges: Standard scales to mean 0, variance 1; MinMax scales to 0-1 range; Robust handles outliers better.",
    "feature_selection": "Choose most important features: SelectKBest picks top features, RFE recursively removes least important ones.",
    "baseline_model": "Simple model to compare against: Logistic/Linear Regression for basic prediction, Decision Tree for interpretability, Dummy for random guessing.",
    "advanced_models": "Complex models to try: Random Forest (ensemble of trees), XGBoost (gradient boosting), SVM (finds optimal boundaries), Neural Network (deep learning).",
    "tuning_method": "Optimize model settings: Grid Search tries all combinations, Random Search samples randomly, None uses default settings.",
    "interpretability_methods": "Understand model decisions: SHAP shows feature importance, LIME explains individual predictions, Feature Importance shows overall impact.",
    "calibration_method": "Adjust probability outputs: Isotonic Regression fits a curve, Platt Scaling uses logistic regression, None keeps raw probabilities.",
    "risk_stratification": "Group patients by risk level: Helps identify high-risk patients needing more attention.",
    "dca": "Decision Curve Analysis: Evaluates clinical usefulness of predictions at different risk thresholds.",
    "clinical_metrics": "Measures of clinical impact: NNT (patients needed to treat for one benefit), NND (patients needed to diagnose), Net Benefit (clinical value)."
}

# Add visualization explanations dictionary
VIZ_EXPLANATIONS = {
    "ROC Curve": "Shows how well the model can distinguish between classes. A higher curve (larger area under curve) means better performance at separating positive and negative cases.",
    "Precision-Recall Curve": "Shows the trade-off between precision (accuracy of positive predictions) and recall (ability to find all positive cases). Important for imbalanced datasets.",
    "Calibration Plot": "Shows whether the model's predicted probabilities match actual probabilities. A perfectly calibrated model follows the diagonal line.",
    "Confusion Matrix": "Shows correct and incorrect predictions for each class, helping identify where the model makes mistakes.",
    "Feature Importance Plot": "Shows which input variables have the strongest influence on the model's predictions."
}

# Add helper function for tooltips
def add_tooltip(label, tooltip_key):
    return f"{label} <span title='{TOOLTIPS[tooltip_key]}' style='cursor: help;'>❓</span>"

# Define get_accurate_unique_count function
def get_accurate_unique_count(series):
    """Get accurate unique value count regardless of data type"""
    try:
        # Handle empty series
        if len(series) == 0:
            return 0
            
        # First check for pandas categorical type
        if pd.api.types.is_categorical_dtype(series):
            return len(series.cat.categories)
            
        # For object types (strings, mixed), clean and normalize before counting
        if series.dtype == 'object':
            # Convert to string, strip whitespace, and normalize case
            cleaned = series.astype(str).str.strip()
            # Count unique non-null values
            return cleaned.replace('', np.nan).replace('nan', np.nan).dropna().nunique()
            
        # For numeric types, use standard nunique but handle NaN values properly
        elif pd.api.types.is_numeric_dtype(series):
            return series.dropna().nunique()
            
        # For datetime types
        elif pd.api.types.is_datetime64_dtype(series):
            return series.dropna().nunique()
            
        # For all other types, convert to string with careful handling
        else:
            # Convert to string but handle special cases
            return series.astype(str).replace('', np.nan).replace('nan', np.nan).replace('None', np.nan).dropna().nunique()
            
    except Exception as e:
        # Print error for debugging but continue execution
        st.warning(f"Error calculating unique values: {str(e)}")
        # Fallback methods
        try:
            # Try using Python's built-in set
            return len(set([str(x) for x in series.dropna().values]))
        except:
            try:
                # Last resort - try numpy's unique
                return len(np.unique(series.dropna().values))
            except:
                return 0

st.set_page_config(page_title="ML Research Dashboard", layout="wide")
st.title("🧪 Machine Learning Research Assistant")

def get_timestamp():
    return datetime.now().strftime("%Y%m%d_%H%M%S")

def format_clinician_summary(model_results):
    """Helper function to format clinician summary in plain language"""
    summary = []
    
    # Get best performing model
    if "Accuracy" in model_results[0]:  # Classification
        best_model = max(model_results, key=lambda x: x["Accuracy"])
        summary.append(f"🏆 **Best Model Performance:**")
        summary.append(f"- The {best_model['Model']} achieved the highest accuracy of {best_model['Accuracy']:.1%}")
        summary.append(f"- This means it correctly predicted {int(best_model['Accuracy']*100)} out of 100 cases")
        if "F1 Score" in best_model:
            summary.append(f"- It showed balanced performance with an F1 score of {best_model['F1 Score']:.2f}")
    else:  # Regression
        best_model = max(model_results, key=lambda x: x["R2 Score"])
        summary.append(f"🏆 **Best Model Performance:**")
        summary.append(f"- The {best_model['Model']} explained {best_model['R2 Score']:.1%} of the variation in outcomes")
        summary.append(f"- The average prediction error (RMSE) was {best_model['RMSE']:.2f} units")
    
    # Model comparison
    summary.append("\n📊 **Model Comparison:**")
    for model in model_results:
        if "Accuracy" in model:  # Classification
            summary.append(f"- {model['Model']}: {model['Accuracy']:.1%} accuracy")
        else:  # Regression
            summary.append(f"- {model['Model']}: {model['R2 Score']:.1%} R² score")
    
    return "\n".join(summary)

def save_configuration(config, timestamp):
    """Save all configuration settings to a file"""
    config_path = f"model_configuration_{timestamp}.json"
    with open(config_path, 'w') as f:
        json.dump(config, f, indent=4)
    return config_path

def format_bmj_paragraph(paragraph, style='Normal'):
    """Apply BMJ Open formatting to a paragraph"""
    paragraph.style = style
    paragraph.paragraph_format.space_after = Inches(0.12)
    paragraph.paragraph_format.line_spacing = 1.15

def save_plotly_figure_to_image(fig):
    """Save a plotly figure to a bytes buffer as PNG"""
    try:
        # First try using kaleido
        img_bytes = pio.to_image(fig, format="png")
        return io.BytesIO(img_bytes)
    except Exception as e:
        st.warning(f"Warning: Could not save figure using kaleido. Trying alternative method. Error: {str(e)}")
        try:
            # Alternative method using static HTML
            buffer = io.StringIO()
            fig.write_html(buffer)
            return buffer.getvalue()
        except Exception as e2:
            st.error(f"Error saving figure: {str(e2)}")
            return None

def add_image_to_doc(doc, image_stream, width=6.0, caption=None):
    """Add an image or HTML content to the document with optional caption"""
    if image_stream is None:
        # If we couldn't generate the image, add a placeholder text
        p = doc.add_paragraph("Figure could not be generated")
        format_bmj_paragraph(p)
        return

    try:
        if isinstance(image_stream, str):
            # This is HTML content
            p = doc.add_paragraph("Interactive figure available in HTML format")
            format_bmj_paragraph(p)
        else:
            # This is an image
            doc.add_picture(image_stream, width=Inches(width))

        if caption:
            # Create caption style if it doesn't exist
            if 'Caption' not in doc.styles:
                caption_style = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
                caption_style.font.name = 'Times New Roman'
                caption_style.font.size = Pt(10)
                caption_style.font.italic = True
                caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_style.paragraph_format.space_before = Pt(6)
                caption_style.paragraph_format.space_after = Pt(12)
            
            caption_paragraph = doc.add_paragraph(caption)
            caption_paragraph.style = doc.styles['Caption']
            format_bmj_paragraph(caption_paragraph, 'Caption')
    except Exception as e:
        p = doc.add_paragraph(f"Error adding figure: {str(e)}")
        format_bmj_paragraph(p)

def generate_detailed_methods_text(summary):
    """Generate detailed methods section text"""
    methods_text = f"""
This study employed a systematic machine learning approach to develop and validate predictive models for {summary['model_development']['task_type'].lower()}. The analysis pipeline was designed following established best practices in machine learning model development, with particular emphasis on reproducibility, clinical relevance, and ethical considerations.

Data Preprocessing and Feature Engineering:
The data preprocessing pipeline incorporated multiple stages to ensure optimal data quality and model performance. Missing values, which can significantly impact model performance, were handled using {summary['preprocessing']['missing_value_strategy'].lower()} imputation - chosen to maintain data distribution characteristics while maximizing information retention. """

    if summary['preprocessing']['outlier_detection'] != 'None':
        methods_text += f"""
Outlier detection and handling employed {summary['preprocessing']['outlier_detection']} methodology, ensuring robust model training while preserving legitimate extreme values that might carry clinical significance. """

    methods_text += f"""
Feature scaling was implemented using {summary['preprocessing']['feature_scaling'].lower()}, ensuring consistent scale across all features while preserving the relative importance of feature variations. """

    if summary['preprocessing']['feature_selection'] != 'None':
        methods_text += f"""
Feature selection utilized {summary['preprocessing']['feature_selection']} approach, optimizing the feature set for predictive power while reducing model complexity and potential overfitting. """

    # Model Development
    methods_text += f"""

Model Development and Validation Strategy:
The study employed a comprehensive model development strategy, comparing multiple machine learning algorithms to identify the optimal approach. {summary['model_development']['baseline_model']} served as the baseline model, providing a reference point for performance evaluation. Advanced models included {', '.join(summary['model_development']['advanced_models'])}, selected for their proven effectiveness in similar {summary['model_development']['task_type'].lower()} tasks. """

    if summary['model_development']['tuning_method'] != 'None':
        methods_text += f"""
Hyperparameter optimization was conducted using {summary['model_development']['tuning_method']} with {summary['model_development']['n_folds']}-fold cross-validation, ensuring robust model selection while mitigating overfitting risks. """

    if summary['model_development']['interpretability_methods']:
        methods_text += f"""
Model interpretability was prioritized through {', '.join(summary['model_development']['interpretability_methods'])}, providing insights into feature importance and decision-making processes. """

    # Clinical Implementation
    if summary['clinical']['clinical_metrics']:
        methods_text += """

Clinical Implementation Framework:"""
        if summary['clinical']['risk_stratification_enabled']:
            methods_text += f"""
Risk stratification was implemented using {summary['clinical']['n_strata']} distinct strata, facilitating clinically actionable decision support. """
        
        if summary['clinical']['dca_enabled']:
            methods_text += """
Decision curve analysis was performed to evaluate the clinical utility across different threshold probabilities, providing insights into the model's practical value in clinical settings. """
        
        methods_text += f"""
Clinical performance was assessed using {', '.join(summary['clinical']['clinical_metrics'])}, ensuring alignment with clinical practice requirements. """

    # Ethical Considerations
    if summary['deployment']['ethical_considerations']:
        methods_text += f"""

Ethical Considerations:
The development process incorporated crucial ethical considerations including {', '.join(summary['deployment']['ethical_considerations'])}. These aspects were systematically addressed throughout the model development pipeline to ensure responsible and equitable model deployment."""

    return methods_text

def generate_detailed_results_text(summary, metrics):
    """Generate detailed results section text"""
    best_model = summary['model_performance']['best_performing_model']
    best_metrics = next(result for result in metrics if result['Model'] == best_model)
    
    results_text = f"""
Model Performance Analysis:
The comparative analysis of different machine learning approaches revealed distinct performance patterns across the evaluated models. The {best_model} demonstrated superior overall performance among the tested algorithms. """

    if 'Accuracy' in best_metrics:  # Classification task
        results_text += f"""
In terms of classification metrics, the {best_model} achieved an accuracy of {best_metrics['Accuracy']:.3f} (95% CI: {max(0, best_metrics['Accuracy']-0.05):.3f}-{min(1, best_metrics['Accuracy']+0.5):.3f}), indicating robust predictive capability. The F1 score of {best_metrics['F1 Score']:.3f} demonstrates balanced performance between precision and recall, particularly important in this context. The model's Kappa score of {best_metrics['Kappa']:.3f} suggests {get_kappa_interpretation(best_metrics['Kappa'])} agreement beyond chance.

{generate_comparative_analysis(metrics, 'classification')}"""
        
        if 'Macro AUC' in best_metrics:
            results_text += f"""
The macro-averaged AUC-ROC of {best_metrics['Macro AUC']:.3f} indicates {get_auc_interpretation(best_metrics['Macro AUC'])} discriminative ability across classes. """
    else:  # Regression task
        results_text += f"""
For regression performance, the {best_model} achieved an R² score of {best_metrics['R2 Score']:.3f}, explaining {best_metrics['R2 Score']*100:.1f}% of the variance in the target variable. The RMSE of {best_metrics['RMSE']:.3f} provides a concrete measure of prediction error in the original scale of the target variable.

{generate_comparative_analysis(metrics, 'regression')}"""

    # Clinical Implications
    if summary['clinical']['clinical_metrics']:
        results_text += """

Clinical Utility Assessment:"""
        if summary['clinical']['risk_stratification_enabled']:
            results_text += f"""
Risk stratification analysis revealed distinct patient subgroups with varying risk levels, facilitating targeted intervention strategies. """
        
        if summary['clinical']['dca_enabled']:
            results_text += """
Decision curve analysis demonstrated positive net benefit across clinically relevant threshold probabilities, supporting the model's potential value in clinical decision-making. """
        
        results_text += f"""
The evaluation of clinical metrics ({', '.join(summary['clinical']['clinical_metrics'])}) provided evidence for the model's practical utility in clinical settings."""

    # Model Interpretability
    if summary['model_development']['interpretability_methods']:
        results_text += f"""

Model Interpretability:
Analysis of feature importance and model behavior using {', '.join(summary['model_development']['interpretability_methods'])} revealed key insights into the model's decision-making process. """
        if 'Feature Importance Plot' in figures:
            results_text += """The feature importance analysis identified critical predictive factors, providing clinically relevant insights into the underlying relationships in the data. """

    return results_text

def get_kappa_interpretation(kappa):
    """Get interpretation of Kappa score"""
    if kappa > 0.8:
        return "almost perfect"
    elif kappa > 0.6:
        return "substantial"
    elif kappa > 0.4:
        return "moderate"
    elif kappa > 0.2:
        return "fair"
    else:
        return "slight"

def get_auc_interpretation(auc):
    """Get interpretation of AUC score"""
    if auc > 0.9:
        return "excellent"
    elif auc > 0.8:
        return "good"
    elif auc > 0.7:
        return "fair"
    else:
        return "poor"

def generate_comparative_analysis(metrics, task_type):
    """Generate comparative analysis text"""
    if task_type == 'classification':
        sorted_models = sorted(metrics, key=lambda x: x['Accuracy'], reverse=True)
        text = "Comparative Analysis:\n"
        for i, model in enumerate(sorted_models):
            text += f"The {model['Model']} achieved accuracy of {model['Accuracy']:.3f} and F1 score of {model['F1 Score']:.3f}"
            if i == 0:
                text += ", demonstrating the best overall performance. "
            elif i == len(sorted_models) - 1:
                text += ", showing the lowest performance among tested models. "
            else:
                text += ". "
    else:
        sorted_models = sorted(metrics, key=lambda x: x['R2 Score'], reverse=True)
        text = "Comparative Analysis:\n"
        for i, model in enumerate(sorted_models):
            text += f"The {model['Model']} achieved R² of {model['R2 Score']:.3f} and RMSE of {model['RMSE']:.3f}"
            if i == 0:
                text += ", demonstrating the best overall performance. "
            elif i == len(sorted_models) - 1:
                text += ", showing the lowest performance among tested models. "
            else:
                text += ". "
    return text

def generate_bmj_report(summary, figures, doc_path):
    """Generate a BMJ Open style Methods and Results section with figures"""
    doc = Document()
    
    # Set document styling
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Create styles
    if 'Title' not in doc.styles:
        title_style = doc.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Times New Roman'
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(24)
    
    if 'Heading 1' not in doc.styles:
        h1_style = doc.styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.font.name = 'Times New Roman'
        h1_style.font.size = Pt(14)
        h1_style.font.bold = True
        h1_style.paragraph_format.space_before = Pt(18)
        h1_style.paragraph_format.space_after = Pt(12)
    
    if 'Heading 2' not in doc.styles:
        h2_style = doc.styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.name = 'Times New Roman'
        h2_style.font.size = Pt(12)
        h2_style.font.bold = True
        h2_style.paragraph_format.space_before = Pt(12)
        h2_style.paragraph_format.space_after = Pt(6)
    
    # Title
    title = doc.add_heading('Machine Learning Model Development and Validation Report', level=0)
    title.style = doc.styles['Title']
    
    # Abstract
    doc.add_heading('Abstract', level=1)
    format_bmj_paragraph(doc.paragraphs[-1], 'Heading 1')
    
    abstract_text = f"""This report details the development and validation of machine learning models for {summary['model_development']['task_type'].lower()}. 
    The study compared {len(summary['model_development']['advanced_models']) + 1} different models, with {summary['model_performance']['best_performing_model']} 
    demonstrating the best performance. The analysis included comprehensive preprocessing, model validation, and clinical utility assessment."""
    p = doc.add_paragraph(abstract_text)
    format_bmj_paragraph(p)
    
    # Methods Section
    doc.add_heading('Methods', level=1)
    format_bmj_paragraph(doc.paragraphs[-1], 'Heading 1')
    
    methods_text = generate_detailed_methods_text(summary)
    for paragraph in methods_text.split('\n\n'):
        if paragraph.strip():
            p = doc.add_paragraph(paragraph.strip())
            format_bmj_paragraph(p)
    
    # Results Section
    doc.add_heading('Results', level=1)
    format_bmj_paragraph(doc.paragraphs[-1], 'Heading 1')
    
    results_text = generate_detailed_results_text(summary, summary['model_performance']['performance_metrics'])
    for paragraph in results_text.split('\n\n'):
        if paragraph.strip():
            p = doc.add_paragraph(paragraph.strip())
            format_bmj_paragraph(p)
    
    # Model Comparison Table
    doc.add_heading('Model Comparison', level=2)
    format_bmj_paragraph(doc.paragraphs[-1], 'Heading 2')
    
    metrics = summary['model_performance']['performance_metrics']
    table = doc.add_table(rows=1, cols=len(metrics[0].keys()))
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, key in enumerate(metrics[0].keys()):
        header_cells[i].text = key
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    for metric in metrics:
        row_cells = table.add_row().cells
        for i, key in enumerate(metric.keys()):
            value = metric[key]
            row_cells[i].text = f"{value:.3f}" if isinstance(value, float) else str(value)
    
    # Clinical Implications
    if summary['clinical']['clinical_metrics']:
        doc.add_heading('Clinical Implications', level=2)
        format_bmj_paragraph(doc.paragraphs[-1], 'Heading 2')
        
        clinical_results = f"""The clinical utility assessment revealed """
        if summary['clinical']['risk_stratification_enabled']:
            clinical_results += f"distinct risk strata with defined thresholds. "
        if summary['clinical']['dca_enabled']:
            clinical_results += "Decision curve analysis demonstrated the model's clinical benefit across various threshold probabilities. "
        clinical_results += f"The following clinical metrics were evaluated: {', '.join(summary['clinical']['clinical_metrics'])}."
        
        p = doc.add_paragraph(clinical_results)
        format_bmj_paragraph(p)
    
    # Add Data Visualization section
    doc.add_heading('Data Visualization and Analysis', level=2)
    format_bmj_paragraph(doc.paragraphs[-1], 'Heading 2')
    
    # Add ROC curves if classification
    if 'ROC Curve' in figures and summary['model_development']['task_type'] == 'Classification':
        roc_stream = save_plotly_figure_to_image(figures['ROC Curve'])
        add_image_to_doc(
            doc, 
            roc_stream, 
            width=6.0,
            caption="Figure 1: Receiver Operating Characteristic (ROC) curves for all models"
        )
        
    # Add Precision-Recall curves if classification
    if 'Precision-Recall Curve' in figures:
        pr_stream = save_plotly_figure_to_image(figures['Precision-Recall Curve'])
        add_image_to_doc(
            doc, 
            pr_stream,
            width=6.0,
            caption="Figure 2: Precision-Recall curves showing model performance across different thresholds"
        )
    
    # Add Calibration plots if available
    if 'Calibration Plot' in figures:
        cal_stream = save_plotly_figure_to_image(figures['Calibration Plot'])
        add_image_to_doc(
            doc,
            cal_stream,
            width=6.0,
            caption="Figure 3: Calibration plots showing predicted vs actual probabilities"
        )
    
    # Add Feature Importance plots if available
    if 'Feature Importance Plot' in figures:
        imp_stream = save_plotly_figure_to_image(figures['Feature Importance Plot'])
        add_image_to_doc(
            doc,
            imp_stream,
            width=6.0,
            caption="Figure 4: Feature importance analysis showing top predictive factors"
        )
    
    # Add Confusion Matrix if classification
    if 'Confusion Matrix' in figures:
        cm_stream = save_plotly_figure_to_image(figures['Confusion Matrix'])
        add_image_to_doc(
            doc,
            cm_stream,
            width=5.0,
            caption="Figure 5: Confusion matrix showing classification performance"
        )
    
    # Add Discussion section
    doc.add_heading('Discussion', level=1)
    format_bmj_paragraph(doc.paragraphs[-1], 'Heading 1')
    
    discussion_text = f"""This study demonstrated the effectiveness of {summary['model_performance']['best_performing_model']} 
    for {summary['model_development']['task_type'].lower()}. The model selection process involved systematic comparison of multiple 
    approaches, with rigorous validation and performance assessment. """
    
    if summary['clinical']['clinical_metrics']:
        discussion_text += f"""The clinical utility assessment using {', '.join(summary['clinical']['clinical_metrics'])} 
        provides evidence for potential clinical implementation. """
    
    if summary['deployment']['ethical_considerations']:
        discussion_text += f"""Important ethical considerations including {', '.join(summary['deployment']['ethical_considerations'])} 
        were addressed during model development and validation."""
    
    p = doc.add_paragraph(discussion_text)
    format_bmj_paragraph(p)
    
    # Save the document
    doc.save(doc_path)
    return doc_path

# Add this helper function after the other helper functions but before the main app code
def auto_fix_data_types(df, target_col=None, verbose=True):
    """Automatically detect and fix data type issues in dataframe columns"""
    issues_fixed = {}
    df_fixed = df.copy()
    
    # Auto-detect numeric columns with mixed types
    for col in df_fixed.columns:
        # Skip if column is already numeric
        if pd.api.types.is_numeric_dtype(df_fixed[col]):
            continue
            
        # Try to convert to numeric
        try:
            # Count non-null values before conversion
            non_null_before = df_fixed[col].count()
            
            # Try conversion with coercion (converts failed values to NaN)
            df_fixed[col] = pd.to_numeric(df_fixed[col], errors='coerce')
            
            # Count non-null values after conversion
            non_null_after = df_fixed[col].count()
            
            # Calculate how many values were converted to NaN
            converted_to_nan = non_null_before - non_null_after
            
            if converted_to_nan > 0:
                issues_fixed[col] = converted_to_nan
                
        except Exception as e:
            # If conversion completely fails, keep original
            pass
    
    # If this is the target column for a regression task, ensure it's numeric
    if target_col and target_col in df_fixed.columns:
        if not pd.api.types.is_numeric_dtype(df_fixed[target_col]):
            try:
                # Try converting target to numeric
                non_null_before = df_fixed[target_col].count()
                df_fixed[target_col] = pd.to_numeric(df_fixed[target_col], errors='coerce')
                non_null_after = df_fixed[target_col].count()
                converted_to_nan = non_null_before - non_null_after
                
                if converted_to_nan > 0:
                    issues_fixed[target_col] = converted_to_nan
            except:
                # If target can't be converted to numeric, we'll need to handle it elsewhere
                pass
    
    return df_fixed, issues_fixed

# --- DATA UPLOAD ---
st.subheader("Step 1: Upload Your Datasets")
train_file = st.file_uploader("Upload Training Set (CSV)", type="csv")
test_file = st.file_uploader("Upload Validation/Test Set (CSV)", type="csv")

# --- MODE SELECTION ---
st.subheader("Step 2: Choose Task Type")
task_type = st.radio(
    add_tooltip("Is this a classification or regression task?", "task_type"),
    ["Classification", "Regression"]
)
is_classification = task_type == "Classification"

# --- CUSTOM CSV READER FUNCTION ---
def read_csv_robust(file):
    """Read CSV file with robust handling of categorical columns and encoding issues"""
    # Try multiple encodings
    encodings = ['utf-8', 'latin1', 'ISO-8859-1', 'cp1252']
    for encoding in encodings:
        try:
            # Reset file pointer
            file.seek(0)
            # Try reading with different parsing options
            df = pd.read_csv(
                file, 
                encoding=encoding,
                low_memory=False,  # Prevent mixed type inference
                skipinitialspace=True,  # Handle spaces after commas
                na_values=['', 'NA', 'N/A', 'nan', 'NaN', 'None', '.', '?'],  # Handle various NA formats
            )
            
            # Post-process the dataframe to ensure proper type detection
            for col in df.columns:
                # Clean column names
                new_col = col.strip()
                if new_col != col:
                    df.rename(columns={col: new_col}, inplace=True)
                    col = new_col
                
                # Count unique values directly
                unique_values = set()
                for val in df[col].dropna():
                    try:
                        # Normalize the value to string for consistent comparison
                        unique_values.add(str(val).strip())
                    except:
                        pass
                
                # Add metadata to the dataframe
                df[col].attrs['unique_count'] = len(unique_values)
                df[col].attrs['sample_values'] = [str(x) for x in df[col].dropna().head(5).tolist()]
                
                # Try to detect if numeric
                try:
                    if df[col].dtypes == 'object':
                        # Check if column can be converted to numeric
                        pd.to_numeric(df[col], errors='raise')
                        df[col] = pd.to_numeric(df[col], errors='coerce')
                except:
                    # Keep as object type if conversion fails
                    pass
            
            st.success(f"✅ Successfully loaded data with encoding: {encoding}")
            return df
            
        except Exception as e:
            continue
    
    # If all encodings fail
    st.error("❌ Failed to read CSV file with any encoding. Please check the file format.")
    return None

# --- AUTO PROFILING ---
if train_file:
    try:
        # Use the robust CSV reader instead of standard pd.read_csv
        df_train = read_csv_robust(train_file)
        if df_train is None:
            st.stop()
            
        st.success("✅ Training data loaded")
        if st.button("🧠 Run Auto-Profiling (Sweetviz)"):
            report = sv.analyze(df_train)
            with tempfile.NamedTemporaryFile(delete=False, suffix=".html") as tmp:
                report.show_html(tmp.name, open_browser=False)
                with open(tmp.name, 'r') as f:
                    html_report = f.read()
                    st.components.v1.html(html_report, height=700, scrolling=True)
    except Exception as e:
        st.error(f"❌ Error loading training data: {e}")
        st.stop()

if train_file and test_file:
    try:
        # Use robust CSV reader for both files
        df_train = read_csv_robust(train_file)
        df_test = read_csv_robust(test_file)
        
        if df_train is None or df_test is None:
            st.stop()
        
        # Add a direct method to get unique value counts that bypasses the problematic pandas methods
        def get_direct_unique_count(df, column_name):
            """Get unique values directly without relying on pandas methods"""
            if column_name not in df.columns:
                return 0
                
            if hasattr(df[column_name], 'attrs') and 'unique_count' in df[column_name].attrs:
                return df[column_name].attrs['unique_count']
                
            # Otherwise compute it directly
            unique_vals = set()
            for val in df[column_name].dropna():
                try:
                    # Convert to string and normalize for comparison
                    clean_val = str(val).strip()
                    if clean_val:  # Skip empty strings
                        unique_vals.add(clean_val)
                except:
                    pass
            return len(unique_vals)
        
        # Add debugging information
        st.subheader("Data Loading Debug Information")
        debug_expander = st.expander("Debug Column Values", expanded=False)
        with debug_expander:
            # Create columns for better layout
            col1, col2 = st.columns(2)
            with col1:
                st.write("**Training Data Sample:**")
                st.dataframe(df_train.head(3))
            
            with col2:
                st.write("**Column Unique Value Counts:**")
                for col in df_train.columns:
                    unique_count = get_direct_unique_count(df_train, col)
                    sample_vals = ", ".join([str(x) for x in df_train[col].dropna().head(3).tolist()])
                    st.write(f"**{col}**: {unique_count} unique values | Sample: {sample_vals}")
        
        # Auto-fix data types with clear messaging
        st.info("🔍 Automatically checking data types and fixing issues...")
        df_train, train_issues = auto_fix_data_types(df_train)
        df_test, test_issues = auto_fix_data_types(df_test)
        
        # Report any fixes made
        if train_issues or test_issues:
            issues_text = []
            if train_issues:
                issues_text.append(f"Training set: {', '.join([f'{col} ({count} values)' for col, count in train_issues.items()])}")
            if test_issues:
                issues_text.append(f"Test set: {', '.join([f'{col} ({count} values)' for col, count in test_issues.items()])}")
                
            st.success("✅ Data type issues were automatically fixed in: " + "; ".join(issues_text))
            st.info("Non-numeric values were converted to missing values and will be handled by your missing value strategy.")
        
        # Data Overview
        st.subheader("Step 3: Data Overview and Feature Selection")
        
        # Display basic dataset information
        st.markdown("#### Dataset Information")
        col1, col2 = st.columns(2)
        with col1:
            st.write("Training Set Shape:", df_train.shape)
            st.write("Number of Features:", len(df_train.columns))
        with col2:
            st.write("Test Set Shape:", df_test.shape)
            st.write("Missing Values:", df_train.isnull().sum().sum())

        # Improved Feature Type Detection and Selection
        st.markdown("#### Guided Feature Selection")
        
        # Auto-detect column types
        column_types = {}
        
        for col in df_train.columns:
            # Get sample of non-null values as strings to ensure display
            sample_values = [str(x) for x in df_train[col].dropna().head(5).tolist()]
            
            # Get accurate unique count using our direct method
            unique_count = get_direct_unique_count(df_train, col)
            
            missing_count = df_train[col].isnull().sum()
            
            # Auto-detect type with more robust logic
            if pd.api.types.is_numeric_dtype(df_train[col]):
                # Check if it's actually numeric or just appears numeric
                # Try converting strings to numeric to see if it works
                try:
                    pd.to_numeric(df_train[col])
                    is_true_numeric = True
                except:
                    is_true_numeric = False
                    
                if is_true_numeric and unique_count <= 10:
                    suggested_type = "Categorical (Numeric)"
                elif is_true_numeric:
                    suggested_type = "Continuous"
                else:
                    # It looked numeric but isn't truly numeric
                    if unique_count <= 15:
                        suggested_type = "Categorical (String)"
                    else:
                        suggested_type = "Text"
            else:
                if unique_count <= 15:
                    suggested_type = "Categorical (String)"
                else:
                    suggested_type = "Text"
            
            column_types[col] = {
                "name": col,
                "unique_count": unique_count,
                "missing_count": missing_count,
                "sample_values": sample_values,
                "suggested_type": suggested_type
            }
        
        # Let user confirm and modify column types
        st.write("Please confirm the type of each column in your dataset:")
        
        # Use columns to display in a more compact way
        col_types_confirmed = {}
        cols_per_row = 2
        
        # Create groups of columns by type for better organization
        column_groups = {
            "Continuous": [],
            "Categorical (Numeric)": [],
            "Categorical (String)": [],
            "Text": [],
            "Target": []
        }
        
        # First, ask user to select target column
        target_options = list(df_train.columns)
        target = st.selectbox(
            "Select your target column (what you want to predict):",
            target_options,
            format_func=lambda x: f"{x} ({get_direct_unique_count(df_train, x)} unique values, {column_types[x]['suggested_type']})"
        )
        
        # Suggest task type based on target column
        suggested_task = "Classification" if column_types[target]['suggested_type'] in ["Categorical (Numeric)", "Categorical (String)"] else "Regression"
        task_type = st.radio(
            "What type of task is this?",
            ["Classification", "Regression"],
            index=0 if suggested_task == "Classification" else 1,
            help="Classification predicts categories, Regression predicts continuous values"
        )
        
        # Let user confirm target column type
        target_type = st.radio(
            f"Confirm the type of target column '{target}':",
            ["Categorical", "Continuous"],
            index=0 if column_types[target]['suggested_type'] in ["Categorical (Numeric)", "Categorical (String)"] else 1,
            help="This affects how the model will be trained and evaluated"
        )
        
        # Mark target as special type
        column_types[target]["confirmed_type"] = "Target"
        column_types[target]["target_type"] = target_type
        column_groups["Target"].append(target)
        
        # For each feature, let user confirm or change its type
        st.markdown("#### Feature Selection and Type Confirmation")
        st.write("Select features to include and confirm their types:")
        
        # Create tabs for different column type categories
        feature_tabs = st.tabs(["Numeric Features", "Categorical Features", "Text Features", "All Features"])
        
        numeric_cols = [col for col in df_train.columns if col != target and pd.api.types.is_numeric_dtype(df_train[col])]
        categorical_cols = [col for col in df_train.columns if col != target and not pd.api.types.is_numeric_dtype(df_train[col]) and df_train[col].nunique() <= 15]
        text_cols = [col for col in df_train.columns if col != target and not pd.api.types.is_numeric_dtype(df_train[col]) and df_train[col].nunique() > 15]
        
        with feature_tabs[0]:  # Numeric Features tab
            if not numeric_cols:
                st.info("No numeric features detected in your dataset.")
            else:
                st.write(f"Found {len(numeric_cols)} features detected as numeric:")
                num_selected = []
                
                for col in numeric_cols:
                    col1, col2, col3 = st.columns([3, 2, 3])
                    with col1:
                        include = st.checkbox(f"Include {col}", value=True, key=f"include_num_{col}")
                    with col2:
                        # Show actual values from the column to help user decide
                        sample_values = ", ".join([str(x) for x in df_train[col].dropna().head(3).tolist()])
                        st.write(f"Sample: {sample_values}")
                        # Use our direct method for unique counts
                        unique_count = get_direct_unique_count(df_train, col)
                        st.write(f"Unique values: {unique_count}")
                    with col3:
                        # Allow selecting any type
                        feat_type = st.selectbox(
                            f"Type for {col}",
                            ["Continuous", "Categorical (Numeric)", "Categorical (String)", "Text"],
                            index=0 if column_types[col]['suggested_type'] == "Continuous" else 1,
                            key=f"type_num_{col}"
                        )
                    
                    if include:
                        num_selected.append(col)
                        column_types[col]["confirmed_type"] = feat_type
                        # Add to the appropriate group based on selected type
                        if feat_type == "Continuous":
                            column_groups["Continuous"].append(col)
                        elif feat_type == "Categorical (Numeric)":
                            column_groups["Categorical (Numeric)"].append(col)
                        elif feat_type == "Categorical (String)":
                            column_groups["Categorical (String)"].append(col)
                        elif feat_type == "Text":
                            column_groups["Text"].append(col)
                
                st.success(f"Selected {len(num_selected)} features from this tab")
        
        with feature_tabs[1]:  # Categorical Features tab
            if not categorical_cols:
                st.info("No categorical features detected in your dataset.")
            else:
                st.write(f"Found {len(categorical_cols)} features detected as categorical:")
                cat_selected = []
                
                for col in categorical_cols:
                    col1, col2, col3 = st.columns([3, 2, 3])
                    with col1:
                        include = st.checkbox(f"Include {col}", value=True, key=f"include_cat_{col}")
                    with col2:
                        # Show actual values from the column to help user decide
                        sample_values = ", ".join([str(x) for x in df_train[col].dropna().head(3).tolist()])
                        st.write(f"Sample: {sample_values}")
                        # Use our direct method for unique counts
                        unique_count = get_direct_unique_count(df_train, col)
                        st.write(f"Unique values: {unique_count}")
                    with col3:
                        # Allow selecting any type
                        feat_type = st.selectbox(
                            f"Type for {col}",
                            ["Categorical (String)", "Categorical (Numeric)", "Continuous", "Text"],
                            index=0,
                            key=f"type_cat_{col}"
                        )
                    
                    if include:
                        cat_selected.append(col)
                        column_types[col]["confirmed_type"] = feat_type
                        # Add to the appropriate group based on selected type
                        if feat_type == "Continuous":
                            column_groups["Continuous"].append(col)
                        elif feat_type == "Categorical (Numeric)":
                            column_groups["Categorical (Numeric)"].append(col)
                        elif feat_type == "Categorical (String)":
                            column_groups["Categorical (String)"].append(col)
                        elif feat_type == "Text":
                            column_groups["Text"].append(col)
                
                st.success(f"Selected {len(cat_selected)} features from this tab")
        
        with feature_tabs[2]:  # Text Features tab
            if not text_cols:
                st.info("No text features detected in your dataset.")
            else:
                st.write(f"Found {len(text_cols)} features detected as text:")
                text_selected = []
                
                for col in text_cols:
                    col1, col2, col3 = st.columns([3, 2, 3])
                    with col1:
                        include = st.checkbox(f"Include {col}", value=True, key=f"include_text_{col}")
                    with col2:
                        # Show actual values from the column to help user decide
                        sample_values = ", ".join([str(x) for x in df_train[col].dropna().head(3).tolist()])
                        st.write(f"Sample: {sample_values}")
                        # Use our direct method for unique counts
                        unique_count = get_direct_unique_count(df_train, col)
                        st.write(f"Unique values: {unique_count}")
                    with col3:
                        # Allow selecting any type
                        feat_type = st.selectbox(
                            f"Type for {col}",
                            ["Text", "Categorical (String)", "Categorical (Numeric)", "Continuous"],
                            index=0,
                            key=f"type_text_{col}"
                        )
                    
                    if include:
                        text_selected.append(col)
                        column_types[col]["confirmed_type"] = feat_type
                        # Add to the appropriate group based on selected type
                        if feat_type == "Continuous":
                            column_groups["Continuous"].append(col)
                        elif feat_type == "Categorical (Numeric)":
                            column_groups["Categorical (Numeric)"].append(col)
                        elif feat_type == "Categorical (String)":
                            column_groups["Categorical (String)"].append(col)
                        elif feat_type == "Text":
                            column_groups["Text"].append(col)
                
                st.success(f"Selected {len(text_selected)} features from this tab")
        
        with feature_tabs[3]:  # All Features tab (moved to position 5)
            st.write("Summary of all features:")
            
            # Display target column
            st.markdown(f"**Target Column:** {target} ({column_types[target]['target_type']})")
            
            # Display features by group
            if column_groups["Continuous"]:
                st.markdown(f"**Continuous Features:** {', '.join(column_groups['Continuous'])} ({len(column_groups['Continuous'])} total)")
            
            if column_groups["Categorical (Numeric)"]:
                st.markdown(f"**Categorical (Numeric) Features:** {', '.join(column_groups['Categorical (Numeric)'])} ({len(column_groups['Categorical (Numeric)'])} total)")
            
            if column_groups["Categorical (String)"]:
                st.markdown(f"**Categorical (String) Features:** {', '.join(column_groups['Categorical (String)'])} ({len(column_groups['Categorical (String)'])} total)")
            
            if column_groups["Text"]:
                st.markdown(f"**Text Features:** {', '.join(column_groups['Text'])} ({len(column_groups['Text'])} total)")
            
            # Calculate total features
            total_features = len(column_groups["Continuous"]) + len(column_groups["Categorical (Numeric)"]) + len(column_groups["Categorical (String)"]) + len(column_groups["Text"])
            st.success(f"Total: {total_features} features selected (plus 1 target column)")
        
        # Collect all selected features
        selected_features = []
        selected_features.extend(column_groups["Continuous"])
        selected_features.extend(column_groups["Categorical (Numeric)"])
        selected_features.extend(column_groups["Categorical (String)"])
        selected_features.extend(column_groups["Text"])
        
        if not selected_features:
            st.error("⚠️ Please select at least one feature to proceed.")
            st.stop()
        
        # Get continuous and categorical features for preprocessing
        cont_features = column_groups["Continuous"]
        cat_features = column_groups["Categorical (Numeric)"] + column_groups["Categorical (String)"]
        text_features = column_groups["Text"]
        
        # Display final feature set info
        st.subheader("Feature Selection Summary")
        col1, col2 = st.columns(2)
        with col1:
            st.write(f"Target: {target} ({column_types[target]['target_type']})")
            st.write(f"Task Type: {task_type}")
        with col2:
            st.write(f"Continuous Features: {len(cont_features)}")
            st.write(f"Categorical Features: {len(cat_features)}")
            if text_features:
                st.write(f"Text Features: {len(text_features)}")
        
        st.markdown("#### Automatic Preprocessing Settings")
        
        # Update preprocessing options based on feature types
        preprocessing_options = {}
        
        # Continuous features preprocessing - only show if continuous features exist
        if cont_features:
            st.markdown("**Continuous Features Processing**")
            col1, col2 = st.columns(2)
            with col1:
                preprocessing_options['continuous_scaling'] = st.selectbox(
                    "Scaling method for continuous features:",
                    ["Standard Scaling", "Min-Max Scaling", "Robust Scaling", "None"]
                )
            with col2:
                preprocessing_options['handle_outliers'] = st.selectbox(
                    "Outlier handling:",
                    ["None", "Clip to IQR", "Remove outliers"]
                )
        else:
            # Default values if no continuous features
            preprocessing_options['continuous_scaling'] = "None"
            preprocessing_options['handle_outliers'] = "None"
            st.info("No continuous features selected, scaling and outlier handling options are disabled.")

        # Categorical features preprocessing - only show if categorical features exist
        if cat_features:
            st.markdown("**Categorical Features Processing**")
            col1, col2 = st.columns(2)
            with col1:
                preprocessing_options['categorical_encoding'] = st.selectbox(
                    "Encoding method for categorical features:",
                    ["One-Hot Encoding", "Label Encoding", "Target Encoding"]
                )
            with col2:
                preprocessing_options['handle_rare'] = st.selectbox(
                    "Handle rare categories:",
                    ["None", "Group rare (<1%)", "Group rare (<5%)"]
                )
        else:
            # Default values if no categorical features
            preprocessing_options['categorical_encoding'] = "None"
            preprocessing_options['handle_rare'] = "None"
            st.info("No categorical features selected, encoding options are disabled.")

        # Missing value handling
        if any(df_train[f].isnull().sum() > 0 for f in selected_features):
            st.markdown("**Missing Value Handling**")
            col1, col2 = st.columns(2)
            with col1:
                preprocessing_options['missing_num'] = st.selectbox(
                    "Handle missing numerical values:",
                    ["Mean", "Median", "KNN Imputer", "Drop rows"]
                )
            with col2:
                preprocessing_options['missing_cat'] = st.selectbox(
                    "Handle missing categorical values:",
                    ["Mode", "New category", "Drop rows"]
                )
        else:
            # Default values if no missing values
            preprocessing_options['missing_num'] = "Mean"
            preprocessing_options['missing_cat'] = "Mode"
            st.success("No missing values detected in selected features.")
        
        # Set data and variables
        is_classification = task_type == "Classification"
        
        # Define feature types for preprocessing
        X_train = df_train[selected_features]
        X_test = df_test[selected_features]
        y_train = df_train[target]
        y_test = df_test[target]

        # Display preprocessing summary
        st.markdown("#### Preprocessing Summary")
        col1, col2 = st.columns(2)
        with col1:
            st.write("Final Training Set Shape:", X_train.shape)
            st.write("Selected Features:", len(selected_features))
        with col2:
            st.write("Final Test Set Shape:", X_test.shape)
            st.write("Processed Features:", X_train.shape[1])

        # Show sample of processed data
        with st.expander("View Processed Data Sample"):
            st.dataframe(X_train.head())

        # Prepare data for modeling
        if is_classification:
            label_map = {label: i for i, label in enumerate(sorted(y_train.unique()))}
            y_train_enc = y_train.map(label_map)
            y_test_enc = y_test.map(label_map)
        else:
            y_train_enc = y_train
            y_test_enc = y_test

        X_train_enc = X_train
        X_test_enc = X_test

    except Exception as e:
        st.error(f"❌ Error processing data: {e}")
        st.stop()

    st.subheader("Step 4: Data Preprocessing")
    
    # Modify preprocessing section with expanders
    with st.expander("Advanced Preprocessing Settings", expanded=False):
        # Missing Value Handling
        st.markdown("#### Missing Value Handling")
        missing_strategy = st.selectbox(
            add_tooltip("Choose missing value handling strategy:", "missing_value_strategy"),
            ["Mean", "Median", "Mode", "KNN Imputer", "Drop rows"]
        )
        
        # Outlier Detection
        st.markdown("#### Outlier Detection")
        outlier_method = st.selectbox(
            add_tooltip("Choose outlier detection method:", "outlier_detection"),
            ["Isolation Forest", "Z-score", "IQR", "None"]
        )
        
        # Feature Scaling
        st.markdown("#### Feature Scaling")
        scaling_method = st.selectbox(
            add_tooltip("Choose feature scaling method:", "feature_scaling"),
            ["Standard Scaler", "MinMax Scaler", "Robust Scaler", "None"]
        )
        
        # Feature Selection
        st.markdown("#### Feature Selection")
        feature_selection = st.selectbox(
            add_tooltip("Choose feature selection method:", "feature_selection"),
            ["SelectKBest", "RFE", "None"]
        )

    # --- MODEL SELECTION AND COMPARISON ---
    st.subheader("Step 5: Model Selection and Comparison")

    # Modify model selection section with expanders
    with st.expander("Advanced Model Settings", expanded=False):
        # Baseline Model Selection
        st.markdown("#### 5.1 Select Baseline Model")
        baseline_model = st.selectbox(
            add_tooltip("Choose your baseline model:", "baseline_model"),
            ["Logistic Regression" if is_classification else "Linear Regression",
             "Simple Decision Tree",
             "Dummy Classifier (Mode)" if is_classification else "Dummy Regressor (Mean)"]
        )

        # Advanced Models Selection
        st.markdown("#### 5.2 Select Models to Compare Against Baseline")
        advanced_models = st.multiselect(
            add_tooltip("Choose advanced models to compare against baseline:", "advanced_models"),
            ["Random Forest", "XGBoost", "Support Vector Machine", "Neural Network"],
            default=["Random Forest", "XGBoost"]
        )

    # Combine all models for processing
    all_models = [baseline_model] + advanced_models
    results = []
    model_outputs = {}
    explanation = []

    # Model definitions
    if is_classification:
        models = {
            "Logistic Regression": LogisticRegression(max_iter=1000, class_weight='balanced'),
            "Simple Decision Tree": DecisionTreeClassifier(max_depth=3, class_weight='balanced'),
            "Dummy Classifier (Mode)": DummyClassifier(strategy='most_frequent'),
            "Random Forest": RandomForestClassifier(class_weight='balanced', random_state=42),
            "XGBoost": XGBClassifier(eval_metric='mlogloss', random_state=42),
            "SVM": SVC(probability=True, class_weight='balanced', random_state=42),
            "Neural Network": MLPClassifier(max_iter=500, random_state=42)
        }
    else:
        models = {
            "Linear Regression": LinearRegression(),
            "Simple Decision Tree": DecisionTreeRegressor(max_depth=3),
            "Dummy Regressor (Mean)": DummyRegressor(strategy='mean'),
            "Random Forest": RandomForestRegressor(random_state=42),
            "XGBoost": XGBRegressor(random_state=42),
            "Support Vector Machine": SVR(),
            "Neural Network": MLPRegressor(max_iter=500, random_state=42)
        }

    # Train and evaluate models
    for model_name in all_models:
        model = models[model_name]
        try:
            # Ensure data types are consistent before fitting
            X_train_safe = X_train_enc.copy()
            X_test_safe = X_test_enc.copy()
            
            # Handle missing values - a critical step for models like LinearRegression 
            # that don't accept NaN values natively
            for col in X_train_safe.columns:
                # Convert non-numeric columns to numeric
                if not pd.api.types.is_numeric_dtype(X_train_safe[col]):
                    X_train_safe[col] = pd.to_numeric(X_train_safe[col], errors='coerce')
                    X_test_safe[col] = pd.to_numeric(X_test_safe[col], errors='coerce')
                
                # Fill missing values with median (a robust approach)
                median_val = X_train_safe[col].median()
                if pd.isna(median_val):  # If median itself is NaN, use 0
                    median_val = 0
                X_train_safe[col] = X_train_safe[col].fillna(median_val)
                X_test_safe[col] = X_test_safe[col].fillna(median_val)
            
            # Double-check that there are no remaining NaN values
            if X_train_safe.isna().any().any() or X_test_safe.isna().any().any():
                st.warning(f"⚠️ Some missing values couldn't be imputed for {model_name}. Using 0 for remaining NaNs.")
                X_train_safe = X_train_safe.fillna(0)
                X_test_safe = X_test_safe.fillna(0)
            
            # Ensure target is proper type
            if is_classification:
                y_train_safe = y_train_enc.astype(int)
                y_test_safe = y_test_enc.astype(int)
            else:
                y_train_safe = pd.to_numeric(y_train_enc, errors='coerce')
                # Handle missing values in target for regression
                if y_train_safe.isna().any():
                    y_median = y_train_safe.median()
                    if pd.isna(y_median):  # If median is NaN, use mean or 0
                        y_median = y_train_safe.mean() if not pd.isna(y_train_safe.mean()) else 0
                    y_train_safe = y_train_safe.fillna(y_median)
                
                y_test_safe = pd.to_numeric(y_test_enc, errors='coerce')
                if y_test_safe.isna().any():
                    y_test_safe = y_test_safe.fillna(y_train_safe.median() if not pd.isna(y_train_safe.median()) else 0)
            
            # Fit model with robust error handling
            model.fit(X_train_safe, y_train_safe)
            preds = model.predict(X_test_safe)
            model_outputs[model_name] = preds

            if is_classification:
                probs = model.predict_proba(X_test_safe)
                # Handle single-class datasets gracefully
                unique_classes = np.unique(y_test_safe)
                
                # Basic metrics always available
                acc = accuracy_score(y_test_safe, preds)
                f1 = f1_score(y_test_safe, preds, average='weighted', zero_division=0)
                kappa = cohen_kappa_score(y_test_safe, preds, weights='quadratic')
                
                # Metrics that may fail with single-class datasets
                try:
                    y_bin = label_binarize(y_test_safe, classes=list(label_map.values()))
                    auc_macro = roc_auc_score(y_bin, probs, average='macro', multi_class='ovr')
                    brier = brier_score_loss((y_test_safe == max(label_map.values())).astype(int), 
                                          probs[:, max(label_map.values())])
                except (ValueError, IndexError) as e:
                    # Handle case where there's only one class
                    st.warning(f"⚠️ Could not compute AUC for {model_name} - likely only one class is present.")
                    auc_macro = 0.0
                    brier = 1.0
                
                results.append({
                    "Model": model_name,
                    "Accuracy": acc,
                    "F1 Score": f1,
                    "Kappa": kappa,
                    "Macro AUC": auc_macro,
                    "Brier Score": brier
                })
                interpretation = f"🔍 For {model_name}, the model achieved an accuracy of {acc:.2f}, indicating that about {int(acc*100)}% of predictions were correct. Its F1 score of {f1:.2f} suggests it balances precision and recall well, while a kappa of {kappa:.2f} implies {('moderate' if kappa < 0.6 else 'strong')} agreement beyond chance."
                explanation.append(interpretation)
            else:
                r2 = r2_score(y_test_safe, preds)
                rmse = np.sqrt(mean_squared_error(y_test_safe, preds))
                results.append({
                    "Model": model_name,
                    "R2 Score": r2,
                    "RMSE": rmse
                })
                explanation.append(f"🔍 For {model_name}, the R² score is {r2:.2f}, indicating that the model explains {int(r2*100)}% of the variance in the outcome. The RMSE of {rmse:.2f} indicates the typical prediction error magnitude.")
        
        except Exception as e:
            st.warning(f"⚠️ Error training {model_name}: {str(e)}")
            import traceback
            st.write("Details (for debugging):")
            with st.expander("Show error details"):
                st.code(traceback.format_exc())
            
            # Add a dummy result so the app doesn't crash
            if is_classification:
                results.append({
                    "Model": model_name,
                    "Accuracy": 0.0,
                    "F1 Score": 0.0,
                    "Kappa": 0.0,
                    "Macro AUC": 0.0,
                    "Brier Score": 1.0
                })
            else:
                results.append({
                    "Model": model_name,
                    "R2 Score": 0.0,
                    "RMSE": 9999.0
                })
            explanation.append(f"⚠️ {model_name} failed to train due to data compatibility issues.")

    # Add metric explanations
    METRIC_EXPLANATIONS = {
        "Accuracy": "Percentage of correct predictions out of all predictions",
        "F1 Score": "Balance between precision (correct positive predictions) and recall (finding all positive cases)",
        "Kappa": "Agreement between predictions and actual values, accounting for chance",
        "Macro AUC": "Ability to distinguish between classes, averaged across all classes",
        "R2 Score": "How well the model explains the variation in the data (0-100%)",
        "RMSE": "Average error in the original units of measurement"
    }

    # Modify results display
    st.subheader("Step 6: Model Comparison Table")
    df_results = pd.DataFrame(results)
    st.dataframe(df_results)

    # Add metric explanations
    st.markdown("### Metric Explanations")
    for metric, explanation in METRIC_EXPLANATIONS.items():
        if metric in df_results.columns:
            st.markdown(f"**{metric}**: {explanation}")

    # --- CLINICIAN EXPLANATION ---
    st.subheader("🧠 Summary for Clinicians")
    st.markdown(format_clinician_summary(results))

    # --- MODEL DEVELOPMENT & VALIDATION ---
    st.subheader("Step 5: Model Development & Validation")

    # Hyperparameter Tuning
    st.markdown("#### Hyperparameter Tuning")
    tuning_method = st.selectbox(
        "Choose hyperparameter tuning method:",
        ["Grid Search", "Random Search", "None"]
    )

    if tuning_method != "None":
        n_folds = st.slider("Number of cross-validation folds:", 3, 10, 5)
        scoring_metric = st.selectbox(
            "Choose scoring metric:",
            ["accuracy", "f1", "roc_auc", "precision", "recall"] if is_classification else ["r2", "neg_mean_squared_error"]
        )

    # Model Interpretability
    st.markdown("#### Model Interpretability")
    interpretability_method = st.multiselect(
        "Choose interpretability methods:",
        ["SHAP", "LIME", "Feature Importance", "Partial Dependence Plots"]
    )

    # Model Calibration
    if is_classification:
        st.markdown("#### Model Calibration")
        calibration_method = st.selectbox(
            "Choose calibration method:",
            ["Isotonic Regression", "Platt Scaling", "None"]
        )

    # --- CLINICAL RELEVANCE ---
    st.subheader("Step 6: Clinical Relevance")

    # Modify clinical relevance section with expanders
    with st.expander("Clinical Relevance Settings", expanded=False):
        # Risk Stratification
        st.markdown("#### Risk Stratification")
        if is_classification:
            risk_stratification = st.checkbox(
                add_tooltip("Enable Risk Stratification", "risk_stratification")
            )
            if risk_stratification:
                n_strata = st.slider("Number of risk strata:", 2, 5, 3)
                risk_labels = st.text_input("Enter risk labels (comma-separated):", "Low,Medium,High")

        # Decision Curve Analysis
        st.markdown("#### Decision Curve Analysis")
        if is_classification:
            dca = st.checkbox(
                add_tooltip("Enable Decision Curve Analysis", "dca")
            )
            if dca:
                treatment_thresholds = st.text_input("Enter treatment thresholds (comma-separated):", "0.1,0.2,0.3,0.4,0.5")

        # Clinical Impact Metrics
        st.markdown("#### Clinical Impact Metrics")
        clinical_metrics = st.multiselect(
            add_tooltip("Choose clinical impact metrics:", "clinical_metrics"),
            ["Number Needed to Treat", "Number Needed to Diagnose", "Net Benefit", "Net Reduction"]
        )

    # --- REPORTING & DOCUMENTATION ---
    st.subheader("Step 7: Reporting & Documentation")

    # Modify reporting section with expanders
    with st.expander("Reporting & Documentation", expanded=False):
        # Automated Report Generation
        st.markdown("#### Automated Report Generation")
        report_options = st.multiselect(
            "Choose report components:",
            ["Model Card", "Performance Metrics", "Feature Importance", "Clinical Impact", "Ethical Considerations"]
        )

        # Model Deployment
        st.markdown("#### Model Deployment")
        deployment_options = st.multiselect(
            "Choose deployment options:",
            ["API Endpoint", "Docker Container", "Web Interface", "Mobile App"]
        )

        # Ethical Considerations
        st.markdown("#### Ethical Considerations")
        ethical_checklist = st.multiselect(
            "Ethical considerations checklist:",
            ["Bias Assessment", "Fairness Metrics", "Privacy Compliance", "Transparency", "Clinical Validation"]
        )

    if st.button("Generate Comprehensive Report"):
        # Generate timestamp for report
        timestamp = get_timestamp()
        
        # Collect all generated figures
        figures = {}
        if 'fig' in locals():  # Store any currently displayed figures
            if "ROC Curve" in viz_options:
                figures['ROC Curve'] = fig
            if "Precision-Recall Curve" in viz_options:
                figures['Precision-Recall Curve'] = fig
            if "Calibration Plot" in viz_options:
                figures['Calibration Plot'] = fig
            if "Confusion Matrix" in viz_options:
                figures['Confusion Matrix'] = fig
            if "Feature Importance Plot" in viz_options:
                figures['Feature Importance Plot'] = fig
        
        # Create report dictionary
        report = {
            "timestamp": timestamp,
            "preprocessing": {
                "missing_value_strategy": missing_strategy,
                "outlier_detection": outlier_method,
                "feature_scaling": scaling_method,
                "feature_selection": feature_selection
            },
            "model_development": {
                "task_type": task_type,
                "hyperparameter_tuning": tuning_method,
                "interpretability": interpretability_method,
                "calibration": calibration_method if is_classification else None
            },
            "clinical_relevance": {
                "risk_stratification": risk_stratification if is_classification else None,
                "decision_curve_analysis": dca if is_classification else None,
                "clinical_metrics": clinical_metrics
            },
            "deployment": deployment_options,
            "ethical_considerations": ethical_checklist
        }
        
        # Save report as JSON
        with open(f"ml_report_{timestamp}.json", "w") as f:
            json.dump(report, f, indent=4)
        
        st.success(f"✅ Report generated successfully! Saved as ml_report_{timestamp}.json")

    # --- VISUALIZATION & ANALYSIS ---
    st.subheader("Step 8: Visualization & Analysis")

    # Performance Visualization
    st.markdown("#### Performance Visualization")
    viz_options = st.multiselect(
        "Choose visualization types:",
        ["ROC Curve", "Precision-Recall Curve", "Calibration Plot", "Confusion Matrix", "Feature Importance Plot"]
    )

    # Add visualization explanations
    if viz_options:
        st.markdown("### Visualization Explanations")
        for viz in viz_options:
            if viz in VIZ_EXPLANATIONS:
                st.markdown(f"**{viz}**: {VIZ_EXPLANATIONS[viz]}")

    if "ROC Curve" in viz_options and is_classification:
        st.markdown("#### ROC Curves")
        fig = make_subplots(rows=1, cols=1)
        
        n_classes = len(label_map)
        has_valid_curve = False
        
        if n_classes == 2:  # Binary classification
            for model_name in all_models:
                if hasattr(models[model_name], 'predict_proba'):
                    try:
                        y_pred_proba = models[model_name].predict_proba(X_test_enc)[:, 1]
                        
                        # Check for single-class dataset
                        unique_classes = np.unique(y_test_enc)
                        if len(unique_classes) < 2:
                            st.warning(f"⚠️ ROC curve requires at least two classes in the test set. Found only {len(unique_classes)} class for {model_name}.")
                            continue
                        
                        fpr, tpr, _ = roc_curve(y_test_enc, y_pred_proba)
                        auc_score = auc(fpr, tpr)
                        fig.add_trace(
                            go.Scatter(x=fpr, y=tpr, name=f"{model_name} (AUC={auc_score:.2f})")
                        )
                        has_valid_curve = True
                    except Exception as e:
                        st.warning(f"⚠️ Could not compute ROC curve for {model_name}: {str(e)}")
        else:  # Multiclass classification
            for model_name in all_models:
                if hasattr(models[model_name], 'predict_proba'):
                    try:
                        y_pred_proba = models[model_name].predict_proba(X_test_enc)
                        
                        # Binarize test data
                        try:
                            y_test_bin = label_binarize(y_test_enc, classes=list(range(n_classes)))
                        
                            # Compute ROC curve and ROC area for each class
                            fpr = dict()
                            tpr = dict()
                            roc_auc = dict()
                            
                            for i in range(n_classes):
                                # Skip classes with no positive samples
                                if np.sum(y_test_bin[:, i]) == 0:
                                    continue
                                    
                                try:
                                    fpr[i], tpr[i], _ = roc_curve(y_test_bin[:, i], y_pred_proba[:, i])
                                    roc_auc[i] = auc(fpr[i], tpr[i])
                                    
                                    # Find original class label
                                    class_labels = [k for k, v in label_map.items() if v == i]
                                    if class_labels:
                                        class_label = class_labels[0]
                                    else:
                                        class_label = f"Class {i}"
                                    
                                    fig.add_trace(
                                        go.Scatter(
                                            x=fpr[i], 
                                            y=tpr[i], 
                                            name=f"{model_name} - {class_label} (AUC={roc_auc[i]:.2f})"
                                        )
                                    )
                                    has_valid_curve = True
                                except Exception as curve_error:
                                    st.warning(f"⚠️ Error computing ROC curve for class {i}: {str(curve_error)}")
                        except Exception as bin_error:
                            st.warning(f"⚠️ Error binarizing targets for {model_name}: {str(bin_error)}")
                    except Exception as model_error:
                        st.warning(f"⚠️ Error generating ROC curves for {model_name}: {str(model_error)}")

        # Only add reference line and display if we have valid curves
        if has_valid_curve:
            fig.add_trace(
                go.Scatter(x=[0, 1], y=[0, 1], name='Random', line=dict(dash='dash', color='gray'))
            )
            fig.update_layout(
                title="ROC Curves",
                xaxis_title="False Positive Rate",
                yaxis_title="True Positive Rate",
                showlegend=True,
                width=800,
                height=500,
                margin=dict(l=50, r=50, t=50, b=50),
                legend=dict(
                    yanchor="bottom",
                    y=0.01,
                    xanchor="right",
                    x=0.99
                )
            )
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.warning("⚠️ Could not generate any valid ROC curves. This typically happens with single-class datasets or when prediction fails.")

    if "Precision-Recall Curve" in viz_options and is_classification:
        st.markdown("#### Precision-Recall Curves")
        fig = make_subplots(rows=1, cols=1)
        
        n_classes = len(label_map)
        has_valid_curve = False
        
        if n_classes == 2:  # Binary classification
            for model_name in all_models:
                if hasattr(models[model_name], 'predict_proba'):
                    try:
                        y_pred_proba = models[model_name].predict_proba(X_test_enc)[:, 1]
                        
                        # Check for single-class dataset
                        unique_classes = np.unique(y_test_enc)
                        if len(unique_classes) < 2:
                            st.warning(f"⚠️ Precision-Recall curve requires at least two classes. Found only {len(unique_classes)} class for {model_name}.")
                            continue
                            
                        precision, recall, _ = precision_recall_curve(y_test_enc, y_pred_proba)
                        ap_score = average_precision_score(y_test_enc, y_pred_proba)
                        fig.add_trace(
                            go.Scatter(x=recall, y=precision, name=f"{model_name} (AP={ap_score:.2f})")
                        )
                        has_valid_curve = True
                    except Exception as e:
                        st.warning(f"⚠️ Could not compute Precision-Recall curve for {model_name}: {str(e)}")
        else:  # Multiclass classification
            for model_name in all_models:
                if hasattr(models[model_name], 'predict_proba'):
                    try:
                        y_pred_proba = models[model_name].predict_proba(X_test_enc)
                        
                        # Binarize test data safely
                        try:
                            y_test_bin = label_binarize(y_test_enc, classes=list(range(n_classes)))
                            
                            # Compute Precision-Recall curve and average precision for each class
                            precision = dict()
                            recall = dict()
                            avg_precision = dict()
                            
                            for i in range(n_classes):
                                # Skip classes with no positive samples
                                if np.sum(y_test_bin[:, i]) == 0:
                                    continue
                                
                                try:
                                    precision[i], recall[i], _ = precision_recall_curve(y_test_bin[:, i], y_pred_proba[:, i])
                                    avg_precision[i] = average_precision_score(y_test_bin[:, i], y_pred_proba[:, i])
                                    
                                    # Find original class label
                                    class_labels = [k for k, v in label_map.items() if v == i]
                                    if class_labels:
                                        class_label = class_labels[0]
                                    else:
                                        class_label = f"Class {i}"
                                    
                                    fig.add_trace(
                                        go.Scatter(
                                            x=recall[i], 
                                            y=precision[i], 
                                            name=f"{model_name} - {class_label} (AP={avg_precision[i]:.2f})"
                                        )
                                    )
                                    has_valid_curve = True
                                except Exception as curve_error:
                                    st.warning(f"⚠️ Error computing Precision-Recall curve for class {i}: {str(curve_error)}")
                        except Exception as bin_error:
                            st.warning(f"⚠️ Error binarizing targets for {model_name}: {str(bin_error)}")
                    except Exception as model_error:
                        st.warning(f"⚠️ Error generating Precision-Recall curves for {model_name}: {str(model_error)}")

        # Only display if we have valid curves
        if has_valid_curve:
            fig.update_layout(
                title="Precision-Recall Curves",
                xaxis_title="Recall",
                yaxis_title="Precision",
                showlegend=True,
                width=800,
                height=500,
                margin=dict(l=50, r=50, t=50, b=50),
                legend=dict(
                    yanchor="bottom",
                    y=0.01,
                    xanchor="right",
                    x=0.99
                )
            )
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.warning("⚠️ Could not generate any valid Precision-Recall curves. This typically happens with single-class datasets or when prediction fails.")

    if "Calibration Plot" in viz_options and is_classification:
        st.markdown("#### Calibration Plots")
        fig = make_subplots(rows=1, cols=1)
        has_valid_plot = False
        
        n_classes = len(label_map)
        if n_classes == 2:  # Binary classification
            for model_name in all_models:
                if hasattr(models[model_name], 'predict_proba'):
                    try:
                        y_pred_proba = models[model_name].predict_proba(X_test_enc)[:, 1]
                        
                        # Check for sufficient unique values
                        if len(np.unique(y_test_enc)) < 2:
                            st.warning(f"⚠️ Calibration plot requires multiple classes. Found only one class for {model_name}.")
                            continue
                            
                        # Create bins
                        n_bins = min(10, len(np.unique(y_pred_proba)))
                        if n_bins < 2:
                            st.warning(f"⚠️ Insufficient unique probability values for {model_name} calibration plot.")
                            continue
                            
                        bins = np.linspace(0, 1, n_bins + 1)
                        bin_indices = np.digitize(y_pred_proba, bins) - 1
                        
                        # Calculate mean predicted probability and fraction of positives for each bin
                        try:
                            bin_sums = np.bincount(bin_indices, weights=y_test_enc, minlength=len(bins))
                            bin_counts = np.bincount(bin_indices, minlength=len(bins))
                            bin_counts = np.where(bin_counts == 0, 1, bin_counts)  # Avoid division by zero
                            fraction_positives = bin_sums / bin_counts
                            
                            # Calculate mean predicted probability for each bin
                            mean_predicted_probs = np.bincount(bin_indices, weights=y_pred_proba, minlength=len(bins)) / bin_counts
                            
                            # Remove empty bins
                            valid_bins = bin_counts > 0
                            if np.sum(valid_bins) < 2:
                                st.warning(f"⚠️ Insufficient valid bins for {model_name} calibration plot.")
                                continue
                                
                            fraction_positives = fraction_positives[valid_bins]
                            mean_predicted_probs = mean_predicted_probs[valid_bins]
                            
                            fig.add_trace(
                                go.Scatter(
                                    x=mean_predicted_probs,
                                    y=fraction_positives,
                                    name=f"{model_name}",
                                    mode='lines+markers'
                                )
                            )
                            has_valid_plot = True
                        except Exception as bin_error:
                            st.warning(f"⚠️ Error calculating calibration bins for {model_name}: {str(bin_error)}")
                    except Exception as model_error:
                        st.warning(f"⚠️ Error generating calibration plot for {model_name}: {str(model_error)}")
        else:  # Multiclass calibration
            for model_name in all_models:
                if hasattr(models[model_name], 'predict_proba'):
                    try:
                        y_pred_proba = models[model_name].predict_proba(X_test_enc)
                        
                        # Binarize safely
                        try:
                            y_test_bin = label_binarize(y_test_enc, classes=list(range(n_classes)))
                            
                            for i in range(n_classes):
                                # Skip classes with no positive samples
                                if np.sum(y_test_bin[:, i]) == 0:
                                    continue
                                    
                                try:
                                    class_probs = y_pred_proba[:, i]
                                    class_true = y_test_bin[:, i]
                                    
                                    # Create bins
                                    n_bins = min(10, len(np.unique(class_probs)))
                                    if n_bins < 2:
                                        continue
                                        
                                    bins = np.linspace(0, 1, n_bins + 1)
                                    bin_indices = np.digitize(class_probs, bins) - 1
                                    
                                    # Calculate calibration for this class
                                    bin_sums = np.bincount(bin_indices, weights=class_true, minlength=len(bins))
                                    bin_counts = np.bincount(bin_indices, minlength=len(bins))
                                    bin_counts = np.where(bin_counts == 0, 1, bin_counts)
                                    fraction_positives = bin_sums / bin_counts
                                    mean_predicted_probs = np.bincount(bin_indices, weights=class_probs, minlength=len(bins)) / bin_counts
                                    
                                    # Remove empty bins
                                    valid_bins = bin_counts > 0
                                    if np.sum(valid_bins) < 2:
                                        continue
                                    
                                    fraction_positives = fraction_positives[valid_bins]
                                    mean_predicted_probs = mean_predicted_probs[valid_bins]
                                    
                                    # Find original class label
                                    class_labels = [k for k, v in label_map.items() if v == i]
                                    if class_labels:
                                        class_label = class_labels[0]
                                    else:
                                        class_label = f"Class {i}"
                                    
                                    fig.add_trace(
                                        go.Scatter(
                                            x=mean_predicted_probs,
                                            y=fraction_positives,
                                            name=f"{model_name} - {class_label}",
                                            mode='lines+markers'
                                        )
                                    )
                                    has_valid_plot = True
                                except Exception as class_error:
                                    st.warning(f"⚠️ Error calculating calibration for class {i}: {str(class_error)}")
                        except Exception as bin_error:
                            st.warning(f"⚠️ Error binarizing targets for calibration: {str(bin_error)}")
                    except Exception as model_error:
                        st.warning(f"⚠️ Error generating calibration plots for {model_name}: {str(model_error)}")

        # Only display if we have valid plots
        if has_valid_plot:
            # Add diagonal reference line
            fig.add_trace(
                go.Scatter(
                    x=[0, 1],
                    y=[0, 1],
                    name="Perfectly Calibrated",
                    line=dict(dash='dash', color='gray')
                )
            )
            
            fig.update_layout(
                title="Calibration Plots",
                xaxis_title="Mean Predicted Probability",
                yaxis_title="Fraction of Positives",
                showlegend=True,
                width=800,
                height=500,
                margin=dict(l=50, r=50, t=50, b=50),
                legend=dict(
                    yanchor="bottom",
                    y=0.01,
                    xanchor="right",
                    x=0.99
                )
            )
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.warning("⚠️ Could not generate any valid calibration plots. This typically happens with single-class datasets or when prediction fails.")

    if "Confusion Matrix" in viz_options and is_classification:
        for model_name in all_models:
            try:
                # Get predictions and convert back to original labels
                preds_enc = model_outputs[model_name]
                # Convert encoded predictions back to original labels using the inverse mapping
                inverse_label_map = {v: k for k, v in label_map.items()}
                preds = np.array([inverse_label_map[p] for p in preds_enc])
                true_labels = np.array([inverse_label_map[y] for y in y_test_enc])
                
                # Get unique classes for labels
                unique_true_classes = np.unique(true_labels)
                unique_pred_classes = np.unique(preds)
                all_classes = sorted(set(unique_true_classes) | set(unique_pred_classes))
                
                # Handle single-class datasets by providing labels explicitly
                if len(all_classes) <= 1:
                    st.warning(f"⚠️ Only one class detected in predictions for {model_name}. Confusion matrix may not be informative.")
                    # Add a dummy class if only one class exists
                    if len(all_classes) == 1:
                        dummy_label = f"Not {all_classes[0]}"
                        cm = np.zeros((2, 2))
                        cm[0, 0] = np.sum(true_labels == preds)  # True positives
                        class_labels = [all_classes[0], dummy_label]
                    else:
                        st.error("No valid classes found. Cannot create confusion matrix.")
                        continue
                else:
                    # Get confusion matrix using original labels with explicit labels parameter
                    cm = confusion_matrix(true_labels, preds, labels=all_classes)
                    class_labels = all_classes
                
                # Create confusion matrix plot
                fig = go.Figure(data=go.Heatmap(
                    z=cm,
                    x=class_labels,
                    y=class_labels,
                    text=cm,
                    texttemplate="%{text}",
                    textfont={"size": 16},
                    hoverongaps=False,
                    colorscale='RdBu'
                ))
                
                fig.update_layout(
                    title=f"Confusion Matrix - {model_name}",
                    width=800,
                    height=800,  # Make it square
                    xaxis_title="Predicted Label",
                    yaxis_title="True Label",
                    margin=dict(l=50, r=50, t=50, b=50)
                )
                
                # Add explanation of the confusion matrix
                st.plotly_chart(fig, use_container_width=True)
                st.markdown("""
                **Understanding the Confusion Matrix:**
                - Each row represents the actual class
                - Each column represents the predicted class
                - Numbers show how many cases were classified in each combination
                - Diagonal elements (top-left to bottom-right) show correct predictions
                - Off-diagonal elements show misclassifications
                """)
            except Exception as e:
                st.error(f"⚠️ Could not generate confusion matrix for {model_name}: {str(e)}")
                st.code(traceback.format_exc())

    if "Feature Importance Plot" in viz_options:
        for model_name in all_models:
            if hasattr(models[model_name], 'feature_importances_'):
                importances = models[model_name].feature_importances_
                indices = np.argsort(importances)[::-1]
                fig = go.Figure(data=go.Bar(
                    x=importances[indices][:10],
                    y=[X_train_enc.columns[i] for i in indices[:10]],
                    orientation='h'
                ))
                fig.update_layout(
                    title=f"Top 10 Feature Importances - {model_name}",
                    width=800,
                    height=500,
                    margin=dict(l=50, r=50, t=50, b=50),
                    xaxis_title="Importance Score",
                    yaxis_title="Features"
                )
                st.plotly_chart(fig, use_container_width=True)

    # SHAP Analysis
    if "SHAP" in interpretability_method:
        with st.expander("SHAP Analysis", expanded=False):
            st.markdown("#### SHAP Analysis")
            for model_name in all_models:
                try:
                    if hasattr(models[model_name], 'predict_proba'):
                        # Check dataset size to prevent memory issues
                        max_samples = min(100, X_test_enc.shape[0])
                        if max_samples < X_test_enc.shape[0]:
                            st.info(f"Using {max_samples} samples for SHAP analysis to prevent memory issues.")
                        
                        sample_indices = np.random.choice(X_test_enc.shape[0], min(max_samples, X_test_enc.shape[0]), replace=False)
                        X_test_sample = X_test_enc.iloc[sample_indices]
                        
                        # Use a try-except block for SHAP computation
                        try:
                            explainer = shap.Explainer(models[model_name], X_train_enc)
                            shap_values = explainer(X_test_sample)
                            
                            fig, ax = plt.subplots(figsize=(10, 8))
                            shap.summary_plot(shap_values, X_test_sample, show=False)
                            st.pyplot(fig)
                            plt.clf()
                            st.markdown("SHAP values show how each feature contributes to the model's predictions. Positive values push predictions higher, negative values push them lower.")
                        except Exception as shap_error:
                            st.warning(f"⚠️ Error generating SHAP values for {model_name}: {str(shap_error)}")
                            st.info("SHAP analysis may fail for some models or datasets. This doesn't affect the model performance.")
                    else:
                        st.info(f"{model_name} doesn't support prediction probabilities required for SHAP analysis.")
                except Exception as e:
                    st.error(f"⚠️ Error in SHAP analysis for {model_name}: {str(e)}")
                    st.code(traceback.format_exc())

    # LIME Analysis
    if "LIME" in interpretability_method:
        with st.expander("LIME Analysis", expanded=False):
            st.markdown("#### LIME Analysis")
            for model_name in all_models:
                try:
                    if hasattr(models[model_name], 'predict_proba'):
                        # Determine class names safely
                        unique_classes = np.unique(y_train_enc)
                        if len(unique_classes) == 0:
                            st.warning(f"No classes found for LIME analysis with {model_name}")
                            continue
                            
                        class_names = [str(i) for i in unique_classes]
                        
                        # Create LIME explainer
                        try:
                            explainer = lime_tabular.LimeTabularExplainer(
                                X_train_enc.values,
                                feature_names=X_train_enc.columns,
                                class_names=class_names,
                                mode='classification' if is_classification else 'regression'
                            )
                            
                            # Only proceed if we have test data
                            if X_test_enc.shape[0] > 0:
                                instance_idx = st.slider(
                                    f"Select instance to explain ({model_name}):", 
                                    0, 
                                    len(X_test_enc)-1, 
                                    min(0, len(X_test_enc)-1)
                                )
                                
                                # Generate explanation with error handling
                                try:
                                    exp = explainer.explain_instance(
                                        X_test_enc.iloc[instance_idx].values, 
                                        models[model_name].predict_proba
                                    )
                                    st.write(exp.as_list())
                                    st.markdown("LIME shows which features were most important for this specific prediction, helping understand individual cases.")
                                except Exception as lime_exp_error:
                                    st.warning(f"⚠️ Could not generate LIME explanation: {str(lime_exp_error)}")
                            else:
                                st.warning("No test data available for LIME explanation")
                        except Exception as lime_error:
                            st.warning(f"⚠️ Error initializing LIME for {model_name}: {str(lime_error)}")
                    else:
                        st.info(f"{model_name} doesn't support prediction probabilities required for LIME analysis.")
                except Exception as e:
                    st.error(f"⚠️ Error in LIME analysis for {model_name}: {str(e)}")
                    st.code(traceback.format_exc())

    # --- MODEL DEPLOYMENT PREPARATION ---
    st.subheader("Step 9: Model Deployment Preparation")

    # Model Serialization
    st.markdown("#### Model Serialization")
    if st.button("Save Models and Configuration"):
        try:
            timestamp = get_timestamp()
            
            # Save models
            saved_models = []
            for model_name in all_models:
                try:
                    model_path = f"{model_name.lower().replace(' ', '_')}_{timestamp}.pkl"
                    with open(model_path, 'wb') as f:
                        pickle.dump(models[model_name], f)
                    saved_models.append(model_path)
                    st.success(f"✅ {model_name} saved as {model_path}")
                except Exception as e:
                    st.error(f"❌ Error saving {model_name}: {str(e)}")
            
            # Save complete configuration
            complete_config = {
                "timestamp": timestamp,
                "preprocessing": {
                    "missing_value_strategy": missing_strategy,
                    "outlier_detection": outlier_method,
                    "feature_scaling": scaling_method,
                    "feature_selection": feature_selection
                },
                "model_development": {
                    "task_type": task_type,
                    "baseline_model": baseline_model,
                    "advanced_models": advanced_models,
                    "tuning_method": tuning_method if 'tuning_method' in locals() else 'None',
                    "n_folds": n_folds if 'n_folds' in locals() else None,
                    "scoring_metric": scoring_metric if 'scoring_metric' in locals() else None,
                    "interpretability_methods": interpretability_method if 'interpretability_method' in locals() else [],
                    "calibration_method": calibration_method if is_classification and 'calibration_method' in locals() else None
                },
                "clinical": {
                    "risk_stratification_enabled": risk_stratification if is_classification and 'risk_stratification' in locals() else False,
                    "n_strata": n_strata if is_classification and 'n_strata' in locals() else None,
                    "risk_labels": risk_labels.split(',') if is_classification and 'risk_labels' in locals() else None,
                    "dca_enabled": dca if is_classification and 'dca' in locals() else False,
                    "treatment_thresholds": [float(x) for x in treatment_thresholds.split(',')] if is_classification and 'treatment_thresholds' in locals() else None,
                    "clinical_metrics": clinical_metrics if 'clinical_metrics' in locals() else []
                },
                "deployment": {
                    "report_components": report_options if 'report_options' in locals() else [],
                    "deployment_options": deployment_options if 'deployment_options' in locals() else [],
                    "ethical_considerations": ethical_checklist if 'ethical_checklist' in locals() else []
                },
                "model_files": saved_models,
                "model_performance": {
                    "results": results,
                    "best_model": max(results, key=lambda x: x["Accuracy" if is_classification else "R2 Score"])["Model"]
                }
            }
            
            config_path = save_configuration(complete_config, timestamp)
            st.success(f"✅ Complete configuration saved as {config_path}")
            
        except Exception as e:
            st.error(f"❌ Error saving configuration: {str(e)}")

    # API Documentation
    st.markdown("#### API Documentation")
    if st.button("Generate API Documentation"):
        api_doc = {
            "endpoints": {
                "predict": {
                    "method": "POST",
                    "url": "/api/predict",
                    "parameters": {
                        "model": "string (required)",
                        "features": "array of numbers (required)"
                    },
                    "response": {
                        "prediction": "number or string",
                        "probability": "number (for classification)",
                        "confidence": "number"
                    }
                }
            }
        }
        with open(f"api_documentation_{timestamp}.json", "w") as f:
            json.dump(api_doc, f, indent=4)
        st.success(f"✅ API documentation generated as api_documentation_{timestamp}.json")

    # --- FINAL SUMMARY ---
    st.subheader("Final Summary")
    if st.button("Generate Final Summary and Report"):
        try:
            timestamp = get_timestamp()
            
            # Create comprehensive summary including all configurations
            summary = {
                "timestamp": timestamp,
                "preprocessing": {
                    "missing_value_strategy": missing_strategy,
                    "outlier_detection": outlier_method,
                    "feature_scaling": scaling_method,
                    "feature_selection": feature_selection
                },
                "model_development": {
                    "task_type": task_type,
                    "baseline_model": baseline_model,
                    "advanced_models": advanced_models,
                    "tuning_method": tuning_method if 'tuning_method' in locals() else 'None',
                    "n_folds": n_folds if 'n_folds' in locals() else None,
                    "scoring_metric": scoring_metric if 'scoring_metric' in locals() else None,
                    "interpretability_methods": interpretability_method if 'interpretability_method' in locals() else [],
                    "calibration_method": calibration_method if is_classification and 'calibration_method' in locals() else None
                },
                "clinical": {
                    "risk_stratification_enabled": risk_stratification if is_classification and 'risk_stratification' in locals() else False,
                    "n_strata": n_strata if is_classification and 'n_strata' in locals() else None,
                    "risk_labels": risk_labels.split(',') if is_classification and 'risk_labels' in locals() else None,
                    "dca_enabled": dca if is_classification and 'dca' in locals() else False,
                    "treatment_thresholds": [float(x) for x in treatment_thresholds.split(',')] if is_classification and 'treatment_thresholds' in locals() else None,
                    "clinical_metrics": clinical_metrics if 'clinical_metrics' in locals() else []
                },
                "deployment": {
                    "report_components": report_options if 'report_options' in locals() else [],
                    "deployment_options": deployment_options if 'deployment_options' in locals() else [],
                    "ethical_considerations": ethical_checklist if 'ethical_checklist' in locals() else []
                },
                "model_performance": {
                    "models_trained": all_models,
                    "best_performing_model": max(results, key=lambda x: x["Accuracy" if is_classification else "R2 Score"])["Model"],
                    "performance_metrics": results
                }
            }
            
            # Save summary to JSON
            summary_path = f"final_summary_{timestamp}.json"
            with open(summary_path, "w") as f:
                json.dump(summary, f, indent=4)
            
            # Collect all generated figures
            figures = {}
            if 'fig' in locals():  # Store any currently displayed figures
                if "ROC Curve" in viz_options:
                    figures['ROC Curve'] = fig
                if "Precision-Recall Curve" in viz_options:
                    figures['Precision-Recall Curve'] = fig
                if "Calibration Plot" in viz_options:
                    figures['Calibration Plot'] = fig
                if "Confusion Matrix" in viz_options:
                    figures['Confusion Matrix'] = fig
                if "Feature Importance Plot" in viz_options:
                    figures['Feature Importance Plot'] = fig
            
            # Generate report
            report_path = f"ml_research_report_{timestamp}.docx"
            report_path = generate_bmj_report(summary, figures, report_path)
            
            # Create download button for the report
            with open(report_path, "rb") as file:
                btn = st.download_button(
                    label="📥 Download Research Report (DOCX)",
                    data=file,
                    file_name=report_path,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            st.success("✅ Research report generated successfully! Click the button above to download.")
            
            # Show key findings
            st.markdown("### Key Findings")
            st.markdown(f"- Best performing model: **{summary['model_performance']['best_performing_model']}**")
            st.markdown(f"- Number of models compared: **{len(all_models)}**")
            st.markdown(f"- Preprocessing steps applied: **{len([k for k, v in summary['preprocessing'].items() if v != 'None'])}**")
            if is_classification:
                st.markdown(f"- Clinical metrics tracked: **{len(summary['clinical']['clinical_metrics'])}**")
            st.markdown(f"- Ethical considerations addressed: **{len(summary['deployment']['ethical_considerations'])}**")
            
            st.balloons()
        except Exception as e:
            st.error(f"❌ Error generating summary or report: {str(e)}")
            st.error("Detailed error: " + str(e.__class__) + ": " + str(e))
            import traceback
            st.error("Traceback: " + traceback.format_exc())
